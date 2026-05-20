from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
OUT = ROOT / "paper_output"
MD = OUT / "final_paper.md"
DOCX = OUT / "final_paper.docx"


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 14 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clean_inline_markdown(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [clean_inline_markdown(part) for part in line.strip().strip("|").split("|")]
        if cells and not all(set(cell) <= {"-", " "} for cell in cells):
            rows.append(cells)
    return rows


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(col_count):
        text = rows[0][i] if i < len(rows[0]) else ""
        set_cell_text(table.rows[0].cells[i], text, bold=True)
        set_cell_shading(table.rows[0].cells[i], "EAF2F8")

    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(col_count):
            set_cell_text(cells[i], row[i] if i < len(row) else "")

    doc.add_paragraph()


def add_image(doc: Document, alt: str, path_text: str) -> None:
    image_path = OUT / path_text
    if not image_path.exists():
        image_path = OUT / path_text.replace("/", "\\")
    jpg_path = OUT / "figures" / "docx_jpg" / (Path(path_text).stem + ".jpg")
    if jpg_path.exists():
        image_path = jpg_path
    if not image_path.exists():
        p = doc.add_paragraph(f"[图像缺失] {alt}: {path_text}")
        p.runs[0].font.color.rgb = RGBColor(180, 0, 0)
        return
    caption = doc.add_paragraph(clean_inline_markdown(alt))
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(9)
        run.italic = True
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(image_path), width=Inches(5.8))


def add_paragraph(doc: Document, text: str) -> None:
    text = clean_inline_markdown(text)
    if not text:
        return
    if re.match(r"^\d+\.\s+", text):
        p = doc.add_paragraph(text, style=None)
        p.paragraph_format.left_indent = Inches(0.18)
    else:
        p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.22)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.space_after = Pt(4)


def build() -> None:
    content = MD.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)):
        styles[style_name].font.name = "黑体"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    lines = content.splitlines()
    table_buffer: list[str] = []
    code_block = False
    math_block = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|"):
            table_buffer.append(stripped)
            continue
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

        if stripped.startswith("```"):
            code_block = not code_block
            continue
        if code_block:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue

        if stripped == "$$":
            math_block = not math_block
            continue
        if math_block:
            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Cambria Math"
            continue

        match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if match:
            add_image(doc, match.group(1), match.group(2))
            continue

        if stripped.startswith("# "):
            title = clean_inline_markdown(stripped[2:])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(18)
            doc.add_paragraph()
        elif stripped.startswith("## "):
            doc.add_heading(clean_inline_markdown(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(clean_inline_markdown(stripped[4:]), level=2)
        elif stripped:
            add_paragraph(doc, stripped)

    if table_buffer:
        add_markdown_table(doc, table_buffer)

    doc.save(DOCX)
    print(f"已生成完整 Word：{DOCX}")


def main() -> int:
    configure_utf8_stdio()
    build()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
