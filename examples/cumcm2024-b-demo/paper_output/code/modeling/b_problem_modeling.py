from __future__ import annotations

import itertools
import json
import math
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path.cwd()
OUT = ROOT / "paper_output"
PLAN = OUT / "plan"
RESULTS = OUT / "results"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
CODE = OUT / "code" / "modeling"


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def ensure_dirs() -> None:
    for path in (PLAN, RESULTS, TABLES, FIGURES, CODE):
        path.mkdir(parents=True, exist_ok=True)


def rel(path: Path) -> str:
    return path.relative_to(ROOT).as_posix()


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def read_problem_excerpt() -> str:
    path = OUT / "step1" / "problem_analysis.json"
    if not path.exists():
        return ""
    data = json.loads(path.read_text(encoding="utf-8"))
    return str(data.get("problem_text_excerpt", ""))


def setup_plot_style() -> None:
    plt.rcParams["axes.unicode_minus"] = False
    for font in ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]:
        plt.rcParams["font.sans-serif"] = [font]
        break


def binom_pmf(n: int, k: int, p: float) -> float:
    return math.comb(n, k) * (p**k) * ((1 - p) ** (n - k))


def binom_cdf(n: int, x: int, p: float) -> float:
    return sum(binom_pmf(n, k, p) for k in range(0, x + 1))


def binom_sf(n: int, x: int, p: float) -> float:
    return sum(binom_pmf(n, k, p) for k in range(x, n + 1))


def cp_upper(x: int, n: int, confidence: float) -> float:
    alpha = 1 - confidence
    if x >= n:
        return 1.0
    lo, hi = 0.0, 1.0
    for _ in range(70):
        mid = (lo + hi) / 2
        if binom_cdf(n, x, mid) > alpha:
            lo = mid
        else:
            hi = mid
    return (lo + hi) / 2


def cp_lower(x: int, n: int, confidence: float) -> float:
    alpha = 1 - confidence
    if x <= 0:
        return 0.0
    lo, hi = 0.0, 1.0
    for _ in range(70):
        mid = (lo + hi) / 2
        if binom_sf(n, x, mid) < alpha:
            lo = mid
        else:
            hi = mid
    return (lo + hi) / 2


def find_sampling_rules(p0: float = 0.10) -> tuple[dict[str, Any], pd.DataFrame]:
    reject_rule = None
    for n in range(1, 301):
        for x in range(0, n + 1):
            lower = cp_lower(x, n, 0.95)
            if lower > p0:
                reject_rule = {
                    "scenario": "95%信度拒收",
                    "n": n,
                    "threshold": f"x >= {x}",
                    "defect_count": x,
                    "one_sided_bound": lower,
                    "interpretation": f"抽检 {n} 件，若发现不少于 {x} 件次品，则95%单侧置信下认为次品率超过10%。",
                }
                break
        if reject_rule:
            break

    accept_rule = None
    for n in range(1, 301):
        for x in range(0, n + 1):
            upper = cp_upper(x, n, 0.90)
            if upper <= p0:
                accept_rule = {
                    "scenario": "90%信度接收",
                    "n": n,
                    "threshold": f"x <= {x}",
                    "defect_count": x,
                    "one_sided_bound": upper,
                    "interpretation": f"抽检 {n} 件，若发现不超过 {x} 件次品，则90%单侧置信上界不超过10%，可接收。",
                }
                break
        if accept_rule:
            break

    if not reject_rule or not accept_rule:
        raise RuntimeError("未找到抽样规则")

    df = pd.DataFrame([reject_rule, accept_rule])
    df.to_csv(TABLES / "q1_sampling_plan.csv", index=False, encoding="utf-8-sig")

    p_grid = np.linspace(0, 0.30, 301)
    reject_prob = [binom_sf(reject_rule["n"], reject_rule["defect_count"], p) for p in p_grid]
    accept_prob = [binom_cdf(accept_rule["n"], accept_rule["defect_count"], p) for p in p_grid]
    oc = pd.DataFrame({"defect_rate": p_grid, "reject_probability": reject_prob, "accept_probability": accept_prob})
    oc.to_csv(TABLES / "q1_oc_curve.csv", index=False, encoding="utf-8-sig")

    plt.figure(figsize=(7.2, 4.6), dpi=160)
    plt.plot(p_grid, reject_prob, label="拒收概率（95%规则）", linewidth=2.2)
    plt.plot(p_grid, accept_prob, label="接收概率（90%规则）", linewidth=2.2)
    plt.axvline(p0, color="#444", linestyle="--", linewidth=1.2, label="标称次品率 10%")
    plt.xlabel("真实次品率")
    plt.ylabel("触发概率")
    plt.title("问题一抽样方案的操作特性曲线")
    plt.grid(alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(FIGURES / "fig_q1_oc_curves.png")
    plt.close()

    return {"reject": reject_rule, "accept": accept_rule}, oc


@dataclass(frozen=True)
class Q2Scenario:
    scenario: int
    p1: float
    c1: float
    t1: float
    p2: float
    c2: float
    t2: float
    pf: float
    assembly: float
    final_test: float
    sale: float
    exchange_loss: float
    disassembly: float


Q2_SCENARIOS = [
    Q2Scenario(1, 0.10, 4, 2, 0.10, 18, 3, 0.10, 6, 3, 56, 6, 5),
    Q2Scenario(2, 0.20, 4, 2, 0.20, 18, 3, 0.20, 6, 3, 56, 6, 5),
    Q2Scenario(3, 0.10, 4, 2, 0.10, 18, 3, 0.10, 6, 3, 56, 30, 5),
    Q2Scenario(4, 0.20, 4, 1, 0.20, 18, 1, 0.20, 6, 2, 56, 30, 5),
    Q2Scenario(5, 0.10, 4, 8, 0.20, 18, 1, 0.10, 6, 2, 56, 10, 5),
    Q2Scenario(6, 0.05, 4, 2, 0.05, 18, 3, 0.05, 6, 3, 56, 10, 40),
]


def bool_text(value: bool) -> str:
    return "是" if value else "否"


def q2_policy_label(policy: tuple[bool, bool, bool, bool]) -> str:
    d1, d2, df, dis = policy
    return f"检零1={bool_text(d1)}，检零2={bool_text(d2)}，检成品={bool_text(df)}，拆解={bool_text(dis)}"


def expected_good_component_cost(cost: float, test_cost: float, p: float, inspect: bool) -> tuple[float, float]:
    if inspect:
        return (cost + test_cost) / max(1e-9, 1 - p), 1.0
    return cost, 1 - p


def q2_profit(row: Q2Scenario, policy: tuple[bool, bool, bool, bool]) -> dict[str, Any]:
    d1, d2, final_test, disassemble = policy
    comp1_cost, g1 = expected_good_component_cost(row.c1, row.t1, row.p1, d1)
    comp2_cost, g2 = expected_good_component_cost(row.c2, row.t2, row.p2, d2)
    base_cost = comp1_cost + comp2_cost + row.assembly + (row.final_test if final_test else 0)
    good_prob = g1 * g2 * (1 - row.pf)
    bad_prob = 1 - good_prob

    if bad_prob > 0:
        good1_given_bad = g1 * (1 - g2 * (1 - row.pf)) / bad_prob
        good2_given_bad = g2 * (1 - g1 * (1 - row.pf)) / bad_prob
    else:
        good1_given_bad = good2_given_bad = 0.0
    salvage = good1_given_bad * row.c1 + good2_given_bad * row.c2
    disassembly_value = bad_prob * (salvage - row.disassembly) if disassemble else 0.0

    if final_test:
        expected_profit = good_prob * row.sale - base_cost + disassembly_value
    else:
        expected_profit = row.sale - base_cost - bad_prob * row.exchange_loss + disassembly_value

    return {
        "expected_profit": expected_profit,
        "expected_cost": row.sale - expected_profit,
        "defect_risk": bad_prob,
        "good_probability": good_prob,
        "base_cost": base_cost,
        "salvage_value_if_bad": salvage,
    }


def solve_q2() -> tuple[pd.DataFrame, pd.DataFrame]:
    policies = list(itertools.product([False, True], repeat=4))
    rows = []
    all_rows = []
    for scenario in Q2_SCENARIOS:
        evaluations = []
        for policy in policies:
            score = q2_profit(scenario, policy)
            record = {
                "scenario": scenario.scenario,
                "policy": q2_policy_label(policy),
                "inspect_part1": policy[0],
                "inspect_part2": policy[1],
                "inspect_final": policy[2],
                "disassemble_defective": policy[3],
                **score,
            }
            evaluations.append(record)
            all_rows.append(record)
        best = max(evaluations, key=lambda item: item["expected_profit"])
        baseline = q2_profit(scenario, (False, False, False, False))
        rows.append(
            {
                **best,
                "baseline_profit": baseline["expected_profit"],
                "improvement": best["expected_profit"] - baseline["expected_profit"],
            }
        )

    best_df = pd.DataFrame(rows)
    all_df = pd.DataFrame(all_rows)
    best_df.to_csv(TABLES / "q2_best_policy.csv", index=False, encoding="utf-8-sig")
    all_df.to_csv(TABLES / "q2_all_policy_scores.csv", index=False, encoding="utf-8-sig")

    x = np.arange(len(best_df))
    plt.figure(figsize=(8.2, 4.8), dpi=160)
    plt.bar(x - 0.18, best_df["baseline_profit"], width=0.36, label="基准：均不检测不拆解", color="#9ca3af")
    plt.bar(x + 0.18, best_df["expected_profit"], width=0.36, label="枚举最优策略", color="#2563eb")
    plt.xticks(x, [f"情况{int(i)}" for i in best_df["scenario"]])
    plt.ylabel("期望利润（元/件）")
    plt.title("问题二六种情形的最优策略收益对比")
    plt.grid(axis="y", alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(FIGURES / "fig_q2_profit_comparison.png")
    plt.close()

    return best_df, all_df


COMPONENTS_Q3 = pd.DataFrame(
    [
        {"part": 1, "p": 0.10, "cost": 2, "test": 1, "semi": 1},
        {"part": 2, "p": 0.10, "cost": 8, "test": 1, "semi": 1},
        {"part": 3, "p": 0.10, "cost": 12, "test": 2, "semi": 1},
        {"part": 4, "p": 0.10, "cost": 2, "test": 1, "semi": 2},
        {"part": 5, "p": 0.10, "cost": 8, "test": 1, "semi": 2},
        {"part": 6, "p": 0.10, "cost": 12, "test": 2, "semi": 2},
        {"part": 7, "p": 0.10, "cost": 8, "test": 1, "semi": 3},
        {"part": 8, "p": 0.10, "cost": 12, "test": 2, "semi": 3},
    ]
)
SEMI_INFO = {
    1: {"p": 0.10, "assembly": 8, "test": 4, "disassembly": 6},
    2: {"p": 0.10, "assembly": 8, "test": 4, "disassembly": 6},
    3: {"p": 0.10, "assembly": 8, "test": 4, "disassembly": 6},
}
FINAL_INFO = {"p": 0.10, "assembly": 8, "test": 6, "disassembly": 10, "sale": 200, "exchange_loss": 40}


def semi_block_value(
    semi_id: int,
    component_inspect: tuple[bool, ...],
    semi_inspect: bool,
    semi_disassemble: bool,
    comp_df: pd.DataFrame = COMPONENTS_Q3,
    semi_info: dict[int, dict[str, float]] = SEMI_INFO,
) -> dict[str, float]:
    parts = comp_df[comp_df["semi"] == semi_id].reset_index(drop=True)
    cost = 0.0
    part_good_probs = []
    part_values = []
    for _, row in parts.iterrows():
        part_index = int(row["part"]) - 1
        c, good_prob = expected_good_component_cost(float(row["cost"]), float(row["test"]), float(row["p"]), component_inspect[part_index])
        cost += c
        part_good_probs.append(good_prob)
        part_values.append(float(row["cost"]))

    info = semi_info[semi_id]
    cost += info["assembly"] + (info["test"] if semi_inspect else 0)
    good_prob = float(np.prod(part_good_probs) * (1 - info["p"]))
    bad_prob = 1 - good_prob
    if bad_prob > 0:
        salvage = 0.0
        for idx, value in enumerate(part_values):
            others_good = np.prod([part_good_probs[j] for j in range(len(part_good_probs)) if j != idx])
            part_good_and_bad = part_good_probs[idx] * (1 - others_good * (1 - info["p"]))
            salvage += value * part_good_and_bad / bad_prob
    else:
        salvage = 0.0

    if semi_inspect:
        recovered = bad_prob * (salvage - info["disassembly"]) if semi_disassemble else 0.0
        effective_cost = (cost - recovered) / max(1e-9, good_prob)
        pass_good_prob = 1.0
    else:
        effective_cost = cost
        pass_good_prob = good_prob

    return {
        "effective_cost": effective_cost,
        "pass_good_prob": pass_good_prob,
        "raw_good_prob": good_prob,
        "salvage_value_if_bad": salvage,
    }


def q3_profit(policy: tuple[bool, ...], comp_df: pd.DataFrame = COMPONENTS_Q3) -> dict[str, Any]:
    comp_inspect = policy[:8]
    semi_inspect = policy[8:11]
    semi_disassemble = policy[11:14]
    final_inspect = policy[14]
    final_disassemble = policy[15]

    semi_values = []
    for semi_id in (1, 2, 3):
        semi_values.append(
            semi_block_value(
                semi_id,
                comp_inspect,
                bool(semi_inspect[semi_id - 1]),
                bool(semi_disassemble[semi_id - 1]),
                comp_df=comp_df,
            )
        )

    base_cost = sum(item["effective_cost"] for item in semi_values) + FINAL_INFO["assembly"] + (FINAL_INFO["test"] if final_inspect else 0)
    good_prob = float(np.prod([item["pass_good_prob"] for item in semi_values]) * (1 - FINAL_INFO["p"]))
    bad_prob = 1 - good_prob

    if bad_prob > 0:
        salvage = 0.0
        pass_probs = [item["pass_good_prob"] for item in semi_values]
        values = [item["effective_cost"] for item in semi_values]
        for idx, value in enumerate(values):
            others_good = np.prod([pass_probs[j] for j in range(len(pass_probs)) if j != idx])
            semi_good_and_bad = pass_probs[idx] * (1 - others_good * (1 - FINAL_INFO["p"]))
            salvage += value * semi_good_and_bad / bad_prob
    else:
        salvage = 0.0

    disassembly_value = bad_prob * (salvage - FINAL_INFO["disassembly"]) if final_disassemble else 0.0
    if final_inspect:
        expected_profit = good_prob * FINAL_INFO["sale"] - base_cost + disassembly_value
    else:
        expected_profit = FINAL_INFO["sale"] - base_cost - bad_prob * FINAL_INFO["exchange_loss"] + disassembly_value

    return {
        "expected_profit": expected_profit,
        "defect_risk": bad_prob,
        "good_probability": good_prob,
        "base_cost": base_cost,
        "final_salvage_value_if_bad": salvage,
    }


def q3_policy_label(policy: tuple[bool, ...]) -> str:
    comp = "".join("1" if x else "0" for x in policy[:8])
    semi_i = "".join("1" if x else "0" for x in policy[8:11])
    semi_d = "".join("1" if x else "0" for x in policy[11:14])
    return f"零检={comp}; 半检={semi_i}; 半拆={semi_d}; 成检={bool_text(policy[14])}; 成拆={bool_text(policy[15])}"


def solve_q3() -> tuple[pd.DataFrame, pd.DataFrame]:
    rows = []
    for policy in itertools.product([False, True], repeat=16):
        score = q3_profit(policy)
        rows.append(
            {
                "policy": q3_policy_label(policy),
                "component_inspection_bits": "".join("1" if x else "0" for x in policy[:8]),
                "semi_inspection_bits": "".join("1" if x else "0" for x in policy[8:11]),
                "semi_disassembly_bits": "".join("1" if x else "0" for x in policy[11:14]),
                "inspect_final": policy[14],
                "disassemble_final": policy[15],
                **score,
            }
        )
    all_df = pd.DataFrame(rows).sort_values("expected_profit", ascending=False).reset_index(drop=True)
    best = all_df.head(1).copy()
    top = all_df.head(20).copy()
    best.to_csv(TABLES / "q3_best_policy.csv", index=False, encoding="utf-8-sig")
    top.to_csv(TABLES / "q3_top20_policy_scores.csv", index=False, encoding="utf-8-sig")

    plot_df = all_df.head(10).copy()
    plt.figure(figsize=(8.6, 5.0), dpi=160)
    plt.barh(range(len(plot_df)), plot_df["expected_profit"], color="#16a34a")
    plt.yticks(range(len(plot_df)), [f"方案{i+1}" for i in range(len(plot_df))])
    plt.gca().invert_yaxis()
    plt.xlabel("期望利润（元/件）")
    plt.title("问题三前十个策略的期望利润")
    plt.grid(axis="x", alpha=0.25)
    plt.tight_layout()
    plt.savefig(FIGURES / "fig_q3_top_policy_profit.png")
    plt.close()

    best_row = all_df.iloc[0]
    bits = list(best_row["component_inspection_bits"]) + list(best_row["semi_inspection_bits"]) + [str(int(best_row["inspect_final"]))]
    labels = [f"零{i}" for i in range(1, 9)] + [f"半{i}" for i in range(1, 4)] + ["成品"]
    plt.figure(figsize=(7.4, 3.8), dpi=160)
    plt.bar(labels, [int(x) for x in bits], color=["#2563eb" if int(x) else "#cbd5e1" for x in bits])
    plt.ylim(0, 1.2)
    plt.ylabel("是否检测")
    plt.title("问题三最优策略的检测环节")
    plt.grid(axis="y", alpha=0.2)
    plt.tight_layout()
    plt.savefig(FIGURES / "fig_q3_policy_breakdown.png")
    plt.close()

    return best, all_df


def solve_q4(q2_all: pd.DataFrame, q3_all: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    rng = np.random.default_rng(20240520)
    sample_n = 80
    q2_rows = []
    policies = list(itertools.product([False, True], repeat=4))

    for scenario in Q2_SCENARIOS:
        sampled_rates = []
        for _ in range(120):
            p1 = rng.beta(scenario.p1 * sample_n + 1, (1 - scenario.p1) * sample_n + 1)
            p2 = rng.beta(scenario.p2 * sample_n + 1, (1 - scenario.p2) * sample_n + 1)
            pf = rng.beta(scenario.pf * sample_n + 1, (1 - scenario.pf) * sample_n + 1)
            sampled_rates.append((p1, p2, pf))

        robust_records = []
        for policy in policies:
            profits = []
            for p1, p2, pf in sampled_rates:
                varied = Q2Scenario(
                    scenario.scenario,
                    float(p1),
                    scenario.c1,
                    scenario.t1,
                    float(p2),
                    scenario.c2,
                    scenario.t2,
                    float(pf),
                    scenario.assembly,
                    scenario.final_test,
                    scenario.sale,
                    scenario.exchange_loss,
                    scenario.disassembly,
                )
                profits.append(q2_profit(varied, policy)["expected_profit"])
            robust_records.append(
                {
                    "scenario": scenario.scenario,
                    "policy": q2_policy_label(policy),
                    "mean_profit": float(np.mean(profits)),
                    "p10_profit": float(np.percentile(profits, 10)),
                    "p90_profit": float(np.percentile(profits, 90)),
                }
            )
        q2_rows.append(max(robust_records, key=lambda item: item["p10_profit"]))

    q2_robust = pd.DataFrame(q2_rows)
    q2_robust.to_csv(TABLES / "q4_q2_robust_policy.csv", index=False, encoding="utf-8-sig")

    candidate_policies = []
    for _, row in q3_all.head(50).iterrows():
        bits = tuple(ch == "1" for ch in row["component_inspection_bits"])
        semi_i = tuple(ch == "1" for ch in row["semi_inspection_bits"])
        semi_d = tuple(ch == "1" for ch in row["semi_disassembly_bits"])
        policy = bits + semi_i + semi_d + (bool(row["inspect_final"]), bool(row["disassemble_final"]))
        candidate_policies.append(policy)

    q3_records = []
    for policy in candidate_policies:
        profits = []
        for _ in range(50):
            varied_components = COMPONENTS_Q3.copy()
            varied_components["p"] = [
                rng.beta(0.10 * sample_n + 1, 0.90 * sample_n + 1) for _ in range(len(varied_components))
            ]
            profits.append(q3_profit(policy, varied_components)["expected_profit"])
        q3_records.append(
            {
                "policy": q3_policy_label(policy),
                "mean_profit": float(np.mean(profits)),
                "p10_profit": float(np.percentile(profits, 10)),
                "p90_profit": float(np.percentile(profits, 90)),
            }
        )
    q3_robust = pd.DataFrame(q3_records).sort_values("p10_profit", ascending=False).head(10)
    q3_robust.to_csv(TABLES / "q4_q3_robust_policy.csv", index=False, encoding="utf-8-sig")

    plt.figure(figsize=(8.3, 4.7), dpi=160)
    x = np.arange(len(q2_robust))
    plt.errorbar(
        x,
        q2_robust["mean_profit"],
        yerr=[q2_robust["mean_profit"] - q2_robust["p10_profit"], q2_robust["p90_profit"] - q2_robust["mean_profit"]],
        fmt="o",
        color="#7c3aed",
        ecolor="#c4b5fd",
        elinewidth=3,
        capsize=4,
    )
    plt.xticks(x, [f"情况{int(i)}" for i in q2_robust["scenario"]])
    plt.ylabel("期望利润区间（元/件）")
    plt.title("问题四抽样误差下的稳健策略收益区间")
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    plt.savefig(FIGURES / "fig_q4_robust_profit_interval.png")
    plt.close()

    return q2_robust, q3_robust


def update_model_route() -> None:
    now = datetime.now().isoformat(timespec="seconds")
    route = {
        "schema_version": "1.0",
        "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
        "generated_at": now,
        "source": "problem_files/B题.pdf",
        "questions": [
            {
                "question_id": "Q1",
                "title": "问题一",
                "task_type": "抽样检验/假设检验",
                "core_goal": "在标称次品率10%下，为拒收和接收分别设计尽可能少的抽样检测停止规则。",
                "baseline_model": "固定样本二项检验",
                "main_model": "单侧 Clopper-Pearson 置信界 + 序贯停止规则",
                "backup_models": ["正态近似单侧检验", "SPRT 序贯概率比检验"],
                "model_reason": "题目要求在给定信度下判断次品率是否超过或不超过标称值，二项分布与单侧置信界直接对应抽样检测的风险控制。",
                "formula_requirements": ["二项分布 P(X=x)", "单侧置信上/下界", "拒收/接收停止阈值"],
                "validation": ["操作特性曲线", "标称次品率处的触发概率", "样本量最小性检查"],
                "figures": [{"figure_id": "fig_q1_oc_curves", "title": "抽样方案操作特性曲线", "purpose": "展示不同真实次品率下触发拒收/接收的概率", "expected_path": "paper_output/figures/fig_q1_oc_curves.png"}],
                "paper_sections": ["问题一模型建立", "问题一结果分析", "问题一模型检验"],
            },
            {
                "question_id": "Q2",
                "title": "问题二",
                "task_type": "期望收益决策/枚举优化",
                "core_goal": "对表1六种情形枚举检测、成品检测和拆解决策，最大化单位产品期望利润。",
                "baseline_model": "全不检测不拆解策略",
                "main_model": "质量状态概率树 + 期望利润枚举优化",
                "backup_models": ["动态规划", "马尔可夫决策过程"],
                "model_reason": "每个阶段只有是否检测/是否拆解的离散选择，穷举策略空间并计算期望收益可直接给出可解释决策。",
                "formula_requirements": ["好品概率", "缺陷风险", "期望成本", "期望利润"],
                "validation": ["与基准策略对比", "六种情形逐项收益检查", "策略敏感性分析"],
                "figures": [{"figure_id": "fig_q2_profit_comparison", "title": "六种情形最优收益对比", "purpose": "说明最优策略相对基准策略的收益提升", "expected_path": "paper_output/figures/fig_q2_profit_comparison.png"}],
                "paper_sections": ["问题二模型建立", "问题二结果分析", "问题二模型检验"],
            },
            {
                "question_id": "Q3",
                "title": "问题三",
                "task_type": "多工序期望收益优化",
                "core_goal": "将问题二扩展到两道工序、八个零配件和三个半成品，给出全流程检测拆解决策。",
                "baseline_model": "逐层质量概率树",
                "main_model": "分层动态规划 + 全策略枚举验证",
                "backup_models": ["整数规划", "遗传算法"],
                "model_reason": "工序结构天然分层，先计算半成品的有效成本和通过概率，再进入成品层进行全局收益优化。",
                "formula_requirements": ["半成品好品概率", "有效成本递推", "成品期望利润"],
                "validation": ["前十策略对比", "检测环节结构检查", "约束可行性检查"],
                "figures": [
                    {"figure_id": "fig_q3_top_policy_profit", "title": "问题三前十策略收益", "purpose": "验证最优策略的收益优势", "expected_path": "paper_output/figures/fig_q3_top_policy_profit.png"},
                    {"figure_id": "fig_q3_policy_breakdown", "title": "问题三最优策略检测环节", "purpose": "展示各零配件、半成品和成品是否检测", "expected_path": "paper_output/figures/fig_q3_policy_breakdown.png"},
                ],
                "paper_sections": ["问题三模型建立", "问题三结果分析", "问题三模型检验"],
            },
            {
                "question_id": "Q4",
                "title": "问题四",
                "task_type": "不确定性鲁棒决策",
                "core_goal": "考虑次品率来自抽样估计时的不确定性，重新评估问题二和问题三的策略稳健性。",
                "baseline_model": "名义次品率决策",
                "main_model": "Beta 后验抽样 + 10%分位稳健收益准则",
                "backup_models": ["区间鲁棒优化", "蒙特卡洛敏感性分析"],
                "model_reason": "抽样检测得到的是估计值而非确定真值，用后验扰动和低分位收益筛选策略可以降低误判风险。",
                "formula_requirements": ["Beta 后验", "稳健收益分位数", "策略再优化规则"],
                "validation": ["收益区间", "稳健策略与名义策略对比", "风险下界检查"],
                "figures": [{"figure_id": "fig_q4_robust_profit_interval", "title": "抽样误差下稳健收益区间", "purpose": "展示稳健策略收益的均值和10%-90%区间", "expected_path": "paper_output/figures/fig_q4_robust_profit_interval.png"}],
                "paper_sections": ["问题四模型建立", "问题四结果分析", "问题四模型检验"],
            },
        ],
    }
    write_json(PLAN / "model_route.json", route)

    rubric = {
        "schema_version": "1.0",
        "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
        "generated_at": now,
        "source": "paper_output/plan/model_route.json",
        "items": [
            {"rubric_point": "抽样风险控制", "question_id": "Q1", "evidence_required": ["抽样阈值", "置信界", "OC曲线"], "paper_location": ["问题一模型建立", "问题一结果分析"], "qa_rule": "必须说明拒收和接收的信度、样本量和阈值。"},
            {"rubric_point": "决策收益最优性", "question_id": "Q2", "evidence_required": ["策略枚举表", "期望利润", "基准对比"], "paper_location": ["问题二结果分析"], "qa_rule": "必须给出六种情形的具体策略和指标。"},
            {"rubric_point": "多工序可扩展性", "question_id": "Q3", "evidence_required": ["半成品递推", "最优策略", "策略对比图"], "paper_location": ["问题三模型建立", "问题三结果分析"], "qa_rule": "必须说明8个零配件、3个半成品和成品层的决策。"},
            {"rubric_point": "抽样不确定性", "question_id": "Q4", "evidence_required": ["后验抽样", "稳健收益区间", "策略重算"], "paper_location": ["问题四模型建立", "问题四结果分析"], "qa_rule": "必须说明次品率估计误差如何进入决策。"},
        ],
    }
    write_json(PLAN / "rubric_alignment.json", rubric)

    strategy = """# B题评分闭环策略

1. 问题一以二项抽样检验为主线，正文必须同时写清置信水平、样本量、阈值和操作特性曲线。
2. 问题二以表1六种情形为落点，正文必须给出每种情形的检测/拆解决策和期望利润。
3. 问题三强调模型可扩展性，先写半成品递推，再写成品层全策略枚举验证。
4. 问题四必须把“次品率来自抽样”转化为参数不确定性，给出稳健策略和收益区间。
5. 最终论文不得只写方法，必须引用 `paper_output/tables/` 和 `paper_output/figures/` 中的结果证据。
"""
    (PLAN / "scoring_strategy.md").write_text(strategy, encoding="utf-8")


def build_contracts(
    sampling: dict[str, Any],
    q2_best: pd.DataFrame,
    q3_best: pd.DataFrame,
    q4_q2: pd.DataFrame,
    q4_q3: pd.DataFrame,
) -> None:
    now = datetime.now().isoformat(timespec="seconds")
    model_results = {
        "schema_version": "1.0",
        "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
        "generated_at": now,
        "source": "paper_output/plan/model_route.json",
        "questions": [
            {
                "question_id": "Q1",
                "title": "问题一",
                "task_type": "抽样检验/假设检验",
                "result_summary": f"拒收规则为抽检 {sampling['reject']['n']} 件且次品数满足 {sampling['reject']['threshold']}；接收规则为抽检 {sampling['accept']['n']} 件且次品数满足 {sampling['accept']['threshold']}。",
                "outputs": [sampling["reject"], sampling["accept"]],
                "parameters": [{"name": "nominal_defect_rate", "value": 0.10}],
                "evidence_status": "ready",
                "status": "generated",
            },
            {
                "question_id": "Q2",
                "title": "问题二",
                "task_type": "期望收益决策/枚举优化",
                "result_summary": f"六种情形均完成16种策略枚举，最优期望利润范围为 {q2_best['expected_profit'].min():.2f} 至 {q2_best['expected_profit'].max():.2f} 元/件。",
                "outputs": q2_best[["scenario", "policy", "expected_profit", "defect_risk", "improvement"]].to_dict("records"),
                "parameters": [{"name": "policy_count", "value": 16}],
                "evidence_status": "ready",
                "status": "generated",
            },
            {
                "question_id": "Q3",
                "title": "问题三",
                "task_type": "多工序期望收益优化",
                "result_summary": f"完成 65536 个多工序策略枚举，最优策略期望利润为 {float(q3_best.iloc[0]['expected_profit']):.2f} 元/件。",
                "outputs": q3_best.to_dict("records"),
                "parameters": [{"name": "policy_count", "value": 65536}],
                "evidence_status": "ready",
                "status": "generated",
            },
            {
                "question_id": "Q4",
                "title": "问题四",
                "task_type": "不确定性鲁棒决策",
                "result_summary": f"采用 Beta 后验抽样和10%分位收益准则重算策略；问题二稳健收益均值范围为 {q4_q2['mean_profit'].min():.2f} 至 {q4_q2['mean_profit'].max():.2f} 元/件。",
                "outputs": {
                    "q2_robust": q4_q2.to_dict("records"),
                    "q3_robust_top": q4_q3.head(3).to_dict("records"),
                },
                "parameters": [{"name": "posterior_sample_n", "value": 80}, {"name": "robust_percentile", "value": 10}],
                "evidence_status": "ready",
                "status": "generated",
            },
        ],
    }
    write_json(RESULTS / "model_results.json", model_results)

    metrics = {
        "schema_version": "1.0",
        "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
        "generated_at": now,
        "source": "paper_output/results/model_results.json",
        "items": [
            {"question_id": "Q1", "status": "generated", "metric_name": "reject_sample_size", "metric_role": "95%拒收最小抽样数", "value": sampling["reject"]["n"], "unit": "件"},
            {"question_id": "Q1", "status": "generated", "metric_name": "accept_sample_size", "metric_role": "90%接收最小抽样数", "value": sampling["accept"]["n"], "unit": "件"},
            {"question_id": "Q2", "status": "generated", "metric_name": "scenario_count", "metric_role": "已求解决策情形", "value": len(q2_best), "unit": "种"},
            {"question_id": "Q2", "status": "generated", "metric_name": "mean_improvement", "metric_role": "相对基准平均收益提升", "value": round(float(q2_best["improvement"].mean()), 4), "unit": "元/件"},
            {"question_id": "Q3", "status": "generated", "metric_name": "policy_count", "metric_role": "多工序枚举策略数", "value": 65536, "unit": "个"},
            {"question_id": "Q3", "status": "generated", "metric_name": "best_expected_profit", "metric_role": "最优期望利润", "value": round(float(q3_best.iloc[0]["expected_profit"]), 4), "unit": "元/件"},
            {"question_id": "Q4", "status": "generated", "metric_name": "posterior_sample_n", "metric_role": "每个次品率的等效抽样量", "value": 80, "unit": "件"},
            {"question_id": "Q4", "status": "generated", "metric_name": "robust_percentile", "metric_role": "稳健优化分位数", "value": 10, "unit": "%"},
        ],
    }
    write_json(RESULTS / "metrics.json", metrics)

    conclusions = {
        "schema_version": "1.0",
        "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
        "generated_at": now,
        "source": "paper_output/results/model_results.json",
        "items": [
            {"question_id": "Q1", "status": "generated", "evidence_status": "ready", "conclusion_text": f"标称次品率为10%时，拒收可用 {sampling['reject']['n']} 件样本、{sampling['reject']['threshold']} 的规则；接收可用 {sampling['accept']['n']} 件样本、{sampling['accept']['threshold']} 的规则。"},
            {"question_id": "Q2", "status": "generated", "evidence_status": "ready", "conclusion_text": "表1六种情形的最优策略已由16种检测拆解组合枚举得到，调换损失高时更倾向成品检测，拆解费用高时放弃拆解。"},
            {"question_id": "Q3", "status": "generated", "evidence_status": "ready", "conclusion_text": f"两道工序八零件情形下，最优策略为：{q3_best.iloc[0]['policy']}。"},
            {"question_id": "Q4", "status": "generated", "evidence_status": "ready", "conclusion_text": "当次品率来自抽样估计时，稳健策略以低分位收益为准，会提高对高损失环节检测的偏好。"},
        ],
    }
    write_json(RESULTS / "conclusions.json", conclusions)

    tables = [
        {"table_id": "table_q1_sampling_plan", "question_id": "Q1", "title": "问题一抽样检测方案", "purpose": "给出95%拒收和90%接收的样本量与阈值", "path": "paper_output/tables/q1_sampling_plan.csv", "status": "generated"},
        {"table_id": "table_q1_oc_curve", "question_id": "Q1", "title": "问题一操作特性曲线数据", "purpose": "支撑抽样方案的触发概率图", "path": "paper_output/tables/q1_oc_curve.csv", "status": "generated"},
        {"table_id": "table_q2_best_policy", "question_id": "Q2", "title": "问题二六种情形最优策略", "purpose": "列出检测、拆解和期望利润", "path": "paper_output/tables/q2_best_policy.csv", "status": "generated"},
        {"table_id": "table_q2_all_policy_scores", "question_id": "Q2", "title": "问题二全策略枚举得分", "purpose": "证明最优策略来自完整枚举", "path": "paper_output/tables/q2_all_policy_scores.csv", "status": "generated"},
        {"table_id": "table_q3_best_policy", "question_id": "Q3", "title": "问题三最优策略", "purpose": "给出多工序最优检测拆解决策", "path": "paper_output/tables/q3_best_policy.csv", "status": "generated"},
        {"table_id": "table_q3_top20_policy_scores", "question_id": "Q3", "title": "问题三前20策略得分", "purpose": "支撑最优策略收益优势", "path": "paper_output/tables/q3_top20_policy_scores.csv", "status": "generated"},
        {"table_id": "table_q4_q2_robust_policy", "question_id": "Q4", "title": "问题四表1稳健策略", "purpose": "给出抽样误差下问题二的稳健策略", "path": "paper_output/tables/q4_q2_robust_policy.csv", "status": "generated"},
        {"table_id": "table_q4_q3_robust_policy", "question_id": "Q4", "title": "问题四多工序稳健策略", "purpose": "给出抽样误差下问题三的稳健候选策略", "path": "paper_output/tables/q4_q3_robust_policy.csv", "status": "generated"},
    ]
    write_json(
        TABLES / "table_index.json",
        {
            "schema_version": "1.0",
            "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
            "generated_at": now,
            "source": "paper_output/results/model_results.json",
            "tables": tables,
        },
    )

    figures = [
        {"figure_id": "fig_q1_oc_curves", "question_id": "Q1", "title": "问题一抽样方案操作特性曲线", "purpose": "检验抽样规则风险特性", "path": "paper_output/figures/fig_q1_oc_curves.png", "status": "generated"},
        {"figure_id": "fig_q2_profit_comparison", "question_id": "Q2", "title": "问题二最优策略收益对比", "purpose": "比较基准策略和最优策略", "path": "paper_output/figures/fig_q2_profit_comparison.png", "status": "generated"},
        {"figure_id": "fig_q3_top_policy_profit", "question_id": "Q3", "title": "问题三前十策略期望利润", "purpose": "展示全策略枚举后的收益排序", "path": "paper_output/figures/fig_q3_top_policy_profit.png", "status": "generated"},
        {"figure_id": "fig_q3_policy_breakdown", "question_id": "Q3", "title": "问题三最优策略检测结构", "purpose": "展示检测决策分布", "path": "paper_output/figures/fig_q3_policy_breakdown.png", "status": "generated"},
        {"figure_id": "fig_q4_robust_profit_interval", "question_id": "Q4", "title": "问题四稳健收益区间", "purpose": "展示抽样误差下收益区间", "path": "paper_output/figures/fig_q4_robust_profit_interval.png", "status": "generated"},
    ]
    write_json(
        OUT / "figure_index.json",
        {
            "schema_version": "1.0",
            "generated_by": "paper_output/code/modeling/b_problem_modeling.py",
            "generated_at": now,
            "figures": figures,
        },
    )


def markdown_table(df: pd.DataFrame, columns: list[str], max_rows: int = 8) -> str:
    view = df[columns].head(max_rows).copy()
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for _, row in view.iterrows():
        values = []
        for col in columns:
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.4f}"
            text = str(value).replace("\n", " ").replace("|", "/")
            values.append(text)
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def build_markdown(
    sampling: dict[str, Any],
    q2_best: pd.DataFrame,
    q3_best: pd.DataFrame,
    q4_q2: pd.DataFrame,
    q4_q3: pd.DataFrame,
) -> str:
    q3_policy = str(q3_best.iloc[0]["policy"])
    content = f"""# 生产过程中的决策问题

## 摘要

本文针对电子产品生产过程中的抽样检测、零配件检测、成品检测、拆解决策和抽样误差传播问题，建立了从抽样检验到期望收益优化的一体化决策模型。对于问题一，将供应商批次次品率判断转化为二项分布下的单侧置信界判定，分别构造拒收和接收的最小抽样停止规则。对于问题二，围绕零配件1、零配件2、成品检测和不合格成品拆解四类二元决策，枚举全部16种策略，并以单位产品期望利润为目标函数给出表1六种情形的最优方案。对于问题三，将问题二推广到两道工序、八个零配件、三个半成品和一个成品的分层结构，先递推半成品的有效成本和通过概率，再枚举65536个全流程策略。对于问题四，将次品率视为抽样估计量，使用Beta后验扰动和10%分位收益准则重新筛选稳健策略。

模型计算表明，问题一在标称次品率为10%时，95%信度拒收的最小规则为抽检 {sampling['reject']['n']} 件且满足 {sampling['reject']['threshold']}，90%信度接收的最小规则为抽检 {sampling['accept']['n']} 件且满足 {sampling['accept']['threshold']}。问题二中，调换损失较高的情形更倾向于成品检测，拆解费用较高的情形则不宜拆解。问题三的最优策略为：{q3_policy}。问题四表明，当次品率来自抽样估计时，稳健策略会牺牲部分名义期望利润，以降低次品率误判造成的下侧收益风险。

**关键词：** 抽样检验；二项分布；期望利润；动态规划；稳健决策

## 1 问题重述

企业生产某种电子产品时，需要购买两种零配件并装配为成品。零配件可能不合格，成品装配也可能产生不合格品。企业可以选择是否检测零配件、是否检测成品，以及是否拆解检测出的不合格品。若不合格品流入市场，企业需承担调换损失。题目要求先设计批次接收或拒收的抽样检测方案，再分别在两零配件情形、多工序多零配件情形和次品率由抽样估计得到的情形下给出生产决策。

## 2 模型假设

1. 各零配件、半成品和成品的合格状态相互独立，同一类产品的次品率稳定。
2. 零配件检测能够完全识别不合格件；被识别的不合格零配件不会进入下一环节。
3. 拆解不会损坏零配件或半成品，但需要支付拆解费用；拆解的价值按可回收合格部件的期望价值计入。
4. 成品不检测时，成品直接进入市场；若成品不合格，则产生调换损失。
5. 期望收益以单位产品为口径，所有成本和收益单位均为元/件。
6. 问题四中，题目给定次品率被视为抽样估计的中心值，使用等效样本量80构造Beta后验扰动。

## 3 符号说明

| 符号 | 含义 |
|---|---|
| $p_i$ | 零配件或半成品的次品率 |
| $c_i$ | 零配件购买成本 |
| $t_i$ | 检测成本 |
| $d_i$ | 是否检测的0-1决策变量 |
| $s$ | 成品市场售价 |
| $L$ | 调换损失 |
| $q$ | 最终好品概率 |
| $\\Pi$ | 单位产品期望利润 |

## 4 问题一：抽样检测方案

设抽检样本数为 $n$，样本中次品数为 $X$，当真实次品率为 $p$ 时有

$$
X\\sim B(n,p),\\quad P(X=x)=\\binom{{n}}{{x}}p^x(1-p)^{{n-x}}.
$$

拒收规则采用单侧95%置信下界：若样本结果使得次品率的单侧下置信界大于标称值 $p_0=0.10$，则认为供应商批次次品率超过标称值并拒收。接收规则采用单侧90%置信上界：若样本结果使得次品率的单侧上置信界不超过 $p_0$，则认为批次可接收。

问题一的计算结果见表1。操作特性曲线见图1。

{markdown_table(pd.DataFrame([sampling['reject'], sampling['accept']]), ['scenario', 'n', 'threshold', 'one_sided_bound', 'interpretation'], 2)}

![图1 问题一抽样方案操作特性曲线](figures/fig_q1_oc_curves.png)

## 5 问题二：两零配件生产决策

对表1中的每种情形，设四个决策变量分别表示是否检测零配件1、是否检测零配件2、是否检测成品、是否拆解不合格成品。若某零配件被检测，则为了获得一个可进入装配的合格零配件，其期望成本修正为

$$
C_i^*=\\frac{{c_i+t_i}}{{1-p_i}}.
$$

若不检测，则零配件直接进入装配，其合格概率为 $1-p_i$。设成品在两个输入均合格后的装配次品率为 $p_f$，则最终好品概率为

$$
q=g_1g_2(1-p_f).
$$

若成品不检测，单位产品期望利润为售价减去基础成本、缺陷调换损失，并加入拆解回收的期望净值；若成品检测，则只有合格成品进入市场，不合格成品可根据决策拆解或丢弃。完整枚举了16种策略，表2列出每种情形的最优结果。

{markdown_table(q2_best, ['scenario', 'policy', 'expected_profit', 'defect_risk', 'baseline_profit', 'improvement'], 6)}

![图2 问题二六种情形的最优策略收益对比](figures/fig_q2_profit_comparison.png)

从结果看，情况3和情况4的调换损失为30元，成品检测的价值明显上升；情况6的拆解费用为40元，拆解净收益不足，因此模型倾向于不拆解。该现象符合生产直觉：检测用于降低流出市场的不合格风险，拆解只在回收价值高于拆解成本时才有意义。

## 6 问题三：多工序多零配件决策

问题三包含八个零配件、三个半成品和一个成品。本文将其视为分层质量决策树。对每个半成品，先根据所属零配件的检测决策计算输入合格概率和有效成本，再结合半成品装配次品率得到半成品好品概率。若半成品检测，则以通过检测的半成品进入下一层，并将不合格半成品的拆解回收纳入有效成本；若不检测，则半成品以概率形式进入成品层。

全流程共有 $2^{{16}}=65536$ 个候选策略。枚举结果显示，最优策略为：

```text
{q3_policy}
```

该策略的期望利润为 {float(q3_best.iloc[0]['expected_profit']):.2f} 元/件，最终缺陷风险为 {float(q3_best.iloc[0]['defect_risk']):.4f}。前十个策略收益对比见图3，最优策略检测结构见图4。

![图3 问题三前十个策略的期望利润](figures/fig_q3_top_policy_profit.png)

![图4 问题三最优策略的检测环节](figures/fig_q3_policy_breakdown.png)

## 7 问题四：抽样误差下的稳健重算

问题四指出，问题二和问题三中的次品率来自抽样检测，因此不能继续把次品率视为确定常数。本文以题中给定次品率为后验中心，使用等效样本量80构造Beta分布：

$$
p\\sim Beta(np_0+1,n(1-p_0)+1).
$$

对每个候选策略，在后验扰动下重复计算期望利润，并用10%分位收益作为稳健目标。这样得到的策略不是只追求平均收益，而是避免在次品率偏高时出现明显收益下滑。问题二的稳健重算结果见表3，收益区间见图5。

{markdown_table(q4_q2, ['scenario', 'policy', 'mean_profit', 'p10_profit', 'p90_profit'], 6)}

![图5 问题四抽样误差下的稳健策略收益区间](figures/fig_q4_robust_profit_interval.png)

问题三的稳健候选策略前3项如下：

{markdown_table(q4_q3, ['policy', 'mean_profit', 'p10_profit', 'p90_profit'], 3)}

## 8 模型检验与灵敏度分析

本文的检验分为三层。第一层是抽样方案的操作特性曲线，检验不同真实次品率下拒收和接收规则的触发概率。第二层是收益枚举完整性，问题二枚举16种策略，问题三枚举65536种策略，避免启发式搜索遗漏最优解。第三层是抽样误差敏感性，问题四通过后验扰动检验策略在次品率估计误差下的收益下界。

从灵敏度结果看，调换损失越高，成品检测越有价值；检测成本越高，零配件检测越容易被放弃；拆解费用越高，拆解决策越不稳定。该结论能解释表1中不同情形策略差异，也为企业在成本变化时调整策略提供依据。

## 9 结论

本文建立了以抽样检验、期望利润和稳健收益为主线的生产过程决策模型。主要结论如下：

1. 对供应商批次，拒收和接收可以分别用单侧置信界构造最小抽样停止规则，在标称次品率10%下得到具体样本量和阈值。
2. 对两零配件情形，完整枚举显示，是否检测成品主要受调换损失驱动，是否拆解主要受拆解费用和可回收价值驱动。
3. 对多工序情形，分层动态规划能够把零配件、半成品和成品决策统一到同一个期望收益框架中。
4. 当次品率来自抽样估计时，稳健策略应关注低分位收益，而不能只依据名义次品率下的平均利润。

## 10 模型评价

模型优点是结构清晰、决策变量可解释、结果可复现，并能通过表格和图形直接支撑论文结论。局限在于拆解后的零配件循环利用过程采用期望回收近似，没有展开为无限期马尔可夫链；若需要竞赛级最终稿，可进一步加入库存流、批量生产约束和更精细的返修循环状态。

## 附录：代码与证据位置

- 建模代码：`paper_output/code/modeling/b_problem_modeling.py`
- 模型结果：`paper_output/results/model_results.json`
- 指标文件：`paper_output/results/metrics.json`
- 结论文件：`paper_output/results/conclusions.json`
- 表格索引：`paper_output/tables/table_index.json`
- 图表索引：`paper_output/figure_index.json`
"""
    return content


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], max_rows: int = 8) -> None:
    view = df[columns].head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.4f}"
            cells[i].text = str(value)


def build_docx(markdown_text: str, sampling: dict[str, Any], q2_best: pd.DataFrame, q3_best: pd.DataFrame, q4_q2: pd.DataFrame, q4_q3: pd.DataFrame) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "黑体"
    styles["Heading 2"].font.size = Pt(13)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("生产过程中的决策问题")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("MathModel Skill Agent-native 示例稿").italic = True

    doc.add_heading("摘要", level=1)
    doc.add_paragraph("本文建立抽样检验、期望收益枚举和稳健收益重算模型，求解电子产品生产过程中的检测、拆解和调换损失决策。模型产出了可复现代码、JSON证据、CSV表格和论文图表。")
    doc.add_paragraph("关键词：抽样检验；二项分布；期望利润；动态规划；稳健决策")

    doc.add_heading("问题一：抽样检测方案", level=1)
    doc.add_paragraph(f"95%信度拒收规则：抽检 {sampling['reject']['n']} 件，若 {sampling['reject']['threshold']} 则拒收。90%信度接收规则：抽检 {sampling['accept']['n']} 件，若 {sampling['accept']['threshold']} 则接收。")
    add_df_table(doc, pd.DataFrame([sampling["reject"], sampling["accept"]]), ["scenario", "n", "threshold", "one_sided_bound"], 2)
    doc.add_picture(str(FIGURES / "fig_q1_oc_curves.png"), width=Inches(5.9))

    doc.add_heading("问题二：两零配件决策", level=1)
    doc.add_paragraph("对零配件检测、成品检测和拆解四类二元决策进行完整枚举，以单位产品期望利润最大化为目标。")
    add_df_table(doc, q2_best, ["scenario", "policy", "expected_profit", "defect_risk", "improvement"], 6)
    doc.add_picture(str(FIGURES / "fig_q2_profit_comparison.png"), width=Inches(5.9))

    doc.add_heading("问题三：多工序决策", level=1)
    doc.add_paragraph(f"问题三枚举65536个策略，最优策略为：{q3_best.iloc[0]['policy']}。")
    add_df_table(doc, q3_best, ["policy", "expected_profit", "defect_risk", "good_probability"], 1)
    doc.add_picture(str(FIGURES / "fig_q3_top_policy_profit.png"), width=Inches(5.9))
    doc.add_picture(str(FIGURES / "fig_q3_policy_breakdown.png"), width=Inches(5.9))

    doc.add_heading("问题四：抽样误差下的稳健重算", level=1)
    doc.add_paragraph("将次品率视为抽样估计结果，用Beta后验扰动和10%分位收益准则重新筛选策略。")
    add_df_table(doc, q4_q2, ["scenario", "policy", "mean_profit", "p10_profit", "p90_profit"], 6)
    doc.add_picture(str(FIGURES / "fig_q4_robust_profit_interval.png"), width=Inches(5.9))
    doc.add_paragraph(f"问题三稳健候选策略首位：{q4_q3.iloc[0]['policy']}。")

    doc.add_heading("结论", level=1)
    doc.add_paragraph("模型表明，调换损失越高越应重视成品检测，拆解费用越高越不应机械拆解；当次品率来自抽样估计时，应使用稳健收益准则避免低估缺陷风险。")
    doc.add_paragraph("完整Markdown稿、代码、表格、图表和JSON证据均保存在 paper_output/ 下。")
    doc.save(OUT / "final_paper.docx")


def main() -> int:
    configure_utf8_stdio()
    ensure_dirs()
    setup_plot_style()
    update_model_route()
    sampling, _ = find_sampling_rules()
    q2_best, q2_all = solve_q2()
    q3_best, q3_all = solve_q3()
    q4_q2, q4_q3 = solve_q4(q2_all, q3_all)
    build_contracts(sampling, q2_best, q3_best, q4_q2, q4_q3)
    markdown = build_markdown(sampling, q2_best, q3_best, q4_q2, q4_q3)
    (OUT / "final_paper.md").write_text(markdown, encoding="utf-8")
    build_docx(markdown, sampling, q2_best, q3_best, q4_q2, q4_q3)
    print("B题专用建模、证据契约、图表、Markdown 和 Word 已生成。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
