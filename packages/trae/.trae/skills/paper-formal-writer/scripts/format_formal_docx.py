import csv
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path.cwd()
OUTPUT_DIR = BASE_DIR / "paper_output"
SOURCE_FILE = OUTPUT_DIR / "final_paper_source.md"
FALLBACK_SOURCE_FILE = OUTPUT_DIR / "final_paper.md"
OUTLINE_FILE = OUTPUT_DIR / "plan" / "paper_outline.json"
FIGURE_INDEX_FILE = OUTPUT_DIR / "figure_index.json"
TABLE_INDEX_FILE = OUTPUT_DIR / "tables" / "table_index.json"
DOCX_FILE = OUTPUT_DIR / "final_paper.docx"
REPORT_MD = OUTPUT_DIR / "format_check_report.md"


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def load_json(path: Path) -> Any:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def rel(path: Path) -> str:
    try:
        return path.relative_to(BASE_DIR).as_posix()
    except ValueError:
        return path.as_posix()


def resolve_path(path_text: str) -> Path:
    normalized = path_text.strip().strip("<>").replace("/", "\\")
    path = Path(normalized)
    if path.is_absolute():
        return path
    return BASE_DIR / path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D9D9D9", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margin(cell, margin_twips: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def apply_run_font(run, font_name: str = "宋体", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)

    for name, font_size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    return text.strip()


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(clean_inline_markdown(text))
    apply_run_font(run, "宋体", 10.5)


def add_center_paragraph(document: Document, text: str, font_name: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(clean_inline_markdown(text))
    apply_run_font(run, font_name, size, bold)


def add_heading(document: Document, text: str, level: int) -> None:
    level = max(1, min(level, 3))
    paragraph = document.add_heading(clean_inline_markdown(text), level=level)
    paragraph.paragraph_format.first_line_indent = None
    for run in paragraph.runs:
        apply_run_font(run, "黑体", {1: 15, 2: 13, 3: 12}[level], True)


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(code.rstrip())
    apply_run_font(run, "Consolas", 8.5)


def read_csv_rows(path: Path, max_rows: int = 18, max_cols: int = 8) -> list[list[str]]:
    if not path.exists():
        return []
    encodings = ("utf-8-sig", "utf-8", "gbk")
    for encoding in encodings:
        try:
            with path.open("r", encoding=encoding, newline="") as handle:
                rows = [[str(cell) for cell in row[:max_cols]] for row in csv.reader(handle)]
            return rows[:max_rows]
        except Exception:
            continue
    return []


def add_table_from_rows(document: Document, rows: list[list[str]], caption: str | None = None) -> None:
    if not rows:
        if caption:
            add_center_paragraph(document, caption, bold=True)
        add_body_paragraph(document, "表格数据文件暂不可读取，正式提交前需检查表格索引和源 CSV 文件。")
        return
    if caption:
        add_center_paragraph(document, caption, bold=True)
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            value = row[col_idx] if col_idx < len(row) else ""
            cell.text = clean_inline_markdown(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            set_cell_margin(cell)
            if row_idx == 0:
                set_cell_shading(cell, "F2F2F2")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    apply_run_font(run, "宋体", 9, row_idx == 0)
    document.add_paragraph()


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", stripped):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    add_table_from_rows(document, rows)


def build_table_lookup(table_index: Any) -> dict[str, dict[str, Any]]:
    lookup = {}
    for item in table_index.get("tables", []) if isinstance(table_index, dict) else []:
        if not isinstance(item, dict):
            continue
        table_id = str(item.get("table_id") or "").strip()
        if table_id:
            lookup[table_id] = item
    return lookup


def build_figure_lookup(figure_index: Any) -> dict[str, dict[str, Any]]:
    lookup = {}
    for item in figure_index.get("figures", []) if isinstance(figure_index, dict) else []:
        if not isinstance(item, dict):
            continue
        figure_id = str(item.get("figure_id") or "").strip()
        if figure_id:
            lookup[figure_id] = item
    return lookup


def add_index_table(document: Document, table_id: str, table_lookup: dict[str, dict[str, Any]]) -> bool:
    item = table_lookup.get(table_id)
    if not item:
        return False
    rows = read_csv_rows(resolve_path(str(item.get("path") or "")))
    caption = item.get("caption") or item.get("title") or table_id
    if not str(caption).startswith("表"):
        caption = f"表 {caption}"
    add_table_from_rows(document, rows, caption=str(caption))
    return True


def add_image(document: Document, path: Path, caption: str | None = None) -> bool:
    if not path.exists():
        add_body_paragraph(document, f"图片文件未找到：{rel(path)}。正式提交前需补齐图像文件。")
        return False
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run()
    try:
        run.add_picture(str(path), width=Cm(14.2))
    except Exception:
        try:
            run.add_picture(str(path), width=Cm(12.8))
        except Exception:
            add_body_paragraph(document, f"图片无法插入：{rel(path)}。")
            return False
    if caption:
        add_center_paragraph(document, caption, bold=True)
    return True


def add_index_figure(document: Document, figure_id: str, figure_lookup: dict[str, dict[str, Any]]) -> bool:
    item = figure_lookup.get(figure_id)
    if not item:
        return False
    caption = item.get("caption") or item.get("title") or figure_id
    if not str(caption).startswith("图"):
        caption = f"图 {caption}"
    return add_image(document, resolve_path(str(item.get("path") or item.get("expected_path") or "")), str(caption))


def source_path() -> Path:
    if SOURCE_FILE.exists():
        return SOURCE_FILE
    return FALLBACK_SOURCE_FILE


def render_markdown(document: Document, text: str, table_lookup: dict[str, dict[str, Any]], figure_lookup: dict[str, dict[str, Any]]) -> dict[str, int]:
    stats = {"headings": 0, "tables": 0, "figures": 0, "code_blocks": 0}
    lines = text.splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    in_formula = False
    formula_lines: list[str] = []

    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines))
                stats["code_blocks"] += 1
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if stripped == "$$":
            if in_formula:
                add_center_paragraph(document, "\n".join(formula_lines), font_name="Cambria Math", size=10.5)
                formula_lines = []
                in_formula = False
            else:
                in_formula = True
                formula_lines = []
            idx += 1
            continue
        if in_formula:
            formula_lines.append(stripped)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        table_marker = re.fullmatch(r"\[\[TABLE:([A-Za-z0-9_\-]+)\]\]", stripped, flags=re.IGNORECASE)
        if table_marker:
            if add_index_table(document, table_marker.group(1), table_lookup):
                stats["tables"] += 1
            idx += 1
            continue

        figure_marker = re.fullmatch(r"\[\[FIGURE:([A-Za-z0-9_\-]+)\]\]", stripped, flags=re.IGNORECASE)
        if figure_marker:
            if add_index_figure(document, figure_marker.group(1), figure_lookup):
                stats["figures"] += 1
            idx += 1
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            caption = image_match.group(1).strip()
            path = resolve_path(image_match.group(2).strip())
            if add_image(document, path, caption or None):
                stats["figures"] += 1
            idx += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = [stripped]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("|") and "|" in lines[idx].strip()[1:]:
                table_lines.append(lines[idx].strip())
                idx += 1
            add_markdown_table(document, table_lines)
            stats["tables"] += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            add_heading(document, heading.group(2), level)
            stats["headings"] += 1
            idx += 1
            continue

        numbered_heading = re.match(r"^((?:\d+\.){0,2}\d+)\s+(.+)$", stripped)
        if numbered_heading and len(stripped) <= 80:
            level = numbered_heading.group(1).count(".") + 1
            add_heading(document, stripped, level)
            stats["headings"] += 1
            idx += 1
            continue

        list_item = re.match(r"^[-*]\s+(.+)$", stripped)
        if list_item:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(-0.25)
            run = paragraph.add_run("• " + clean_inline_markdown(list_item.group(1)))
            apply_run_font(run, "宋体", 10.5)
            idx += 1
            continue

        add_body_paragraph(document, stripped)
        idx += 1

    if code_lines:
        add_code_block(document, "\n".join(code_lines))
        stats["code_blocks"] += 1
    if formula_lines:
        add_center_paragraph(document, "\n".join(formula_lines), font_name="Cambria Math", size=10.5)
    return stats


def write_report(stats: dict[str, int], source: Path, outline: Any) -> None:
    lines = [
        "# Formal DOCX Formatting Report",
        "",
        f"- Status: `GENERATED`",
        f"- Generated at: `{datetime.now().isoformat(timespec='seconds')}`",
        f"- Source: `{rel(source)}`",
        f"- Output: `{rel(DOCX_FILE)}`",
        f"- Outline: `{rel(OUTLINE_FILE) if OUTLINE_FILE.exists() else 'missing'}`",
        f"- Title: `{outline.get('title', '') if isinstance(outline, dict) else ''}`",
        f"- Headings: `{stats.get('headings', 0)}`",
        f"- Tables inserted: `{stats.get('tables', 0)}`",
        f"- Figures inserted: `{stats.get('figures', 0)}`",
        f"- Code blocks: `{stats.get('code_blocks', 0)}`",
        "- Render QA: `render_skipped`",
        "",
        "LibreOffice 渲染不是本脚本的强依赖；若本机 LibreOffice 可用，可在最终交付前另行渲染 PNG/PDF 做视觉检查。",
    ]
    REPORT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> int:
    configure_utf8_stdio()
    source = source_path()
    if not source.exists():
        print(f"缺少正式论文 Markdown：{rel(SOURCE_FILE)}", file=sys.stderr)
        return 1

    outline = load_json(OUTLINE_FILE)
    table_index = load_json(TABLE_INDEX_FILE)
    figure_index = load_json(FIGURE_INDEX_FILE)

    document = Document()
    configure_document(document)

    text = source.read_text(encoding="utf-8")
    stats = render_markdown(document, text, build_table_lookup(table_index), build_figure_lookup(figure_index))
    DOCX_FILE.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_FILE)
    write_report(stats, source, outline)
    print(f"正式 Word 已生成：{rel(DOCX_FILE)}")
    print(f"格式化报告已生成：{rel(REPORT_MD)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
