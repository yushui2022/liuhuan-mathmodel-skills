import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


BASE_DIR = Path.cwd()
OUTPUT_DIR = BASE_DIR / "paper_output"
SOURCE_FILE = OUTPUT_DIR / "final_paper_source.md"
FALLBACK_SOURCE_FILE = OUTPUT_DIR / "final_paper.md"
DOCX_FILE = OUTPUT_DIR / "final_paper.docx"
OUTLINE_FILE = OUTPUT_DIR / "plan" / "paper_outline.json"
FIGURE_INDEX_FILE = OUTPUT_DIR / "figure_index.json"
TABLE_INDEX_FILE = OUTPUT_DIR / "tables" / "table_index.json"
REPORT_MD = OUTPUT_DIR / "format_check_report.md"
REPORT_JSON = OUTPUT_DIR / "format_check_report.json"

PLACEHOLDERS = [
    "内容生成中",
    "关键词1",
    "论文题目缺失",
    "TODO",
    "待补",
    "{{",
    "}}",
]

REQUIRED_SECTIONS = [
    "摘要",
    "关键词",
    "1 问题重述",
    "2 问题分析",
    "3 模型假设",
    "4 符号说明",
    "5 模型的建立与求解",
    "6 模型检验",
    "7 模型评价",
    "8 参考文献",
    "附录",
]


def configure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8")
        except Exception:
            pass


def load_json(path: Path) -> Any:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        return {"__error__": str(exc)}


def rel(path: Path) -> str:
    try:
        return path.relative_to(BASE_DIR).as_posix()
    except ValueError:
        return path.as_posix()


def source_path() -> Path:
    if SOURCE_FILE.exists():
        return SOURCE_FILE
    return FALLBACK_SOURCE_FILE


def compact_text(text: str) -> str:
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\[[^\]]+\]\([^)]+\)", "", text)
    text = re.sub(r"`{1,3}", "", text)
    text = re.sub(r"[#>*_\-|$`{}\[\]():;,.，。；：！？、\s]", "", text)
    return text


def char_count(text: str) -> dict[str, int]:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    nonspace = len(re.sub(r"\s+", "", text))
    content = len(compact_text(text))
    return {"cjk": cjk, "nonspace": nonspace, "content": content}


def has_required_section(text: str, label: str) -> bool:
    if label in ("摘要", "关键词", "附录"):
        return re.search(rf"(^|\n)\s*(?:#+\s*)?(?:\*\*)?{re.escape(label)}", text) is not None
    number, title = label.split(" ", 1)
    return re.search(rf"(^|\n)\s*#*\s*{re.escape(number)}\s+.*{re.escape(title)}", text) is not None


def natural_q_key(qid: str) -> tuple[int, str]:
    match = re.search(r"\d+", qid)
    if match:
        return (int(match.group()), qid)
    return (10_000, qid)


def qids_from_outline(outline: Any) -> list[str]:
    if isinstance(outline, dict) and isinstance(outline.get("questions"), list):
        qids = [str(item.get("question_id") or "").strip() for item in outline["questions"] if isinstance(item, dict)]
        qids = [qid for qid in qids if qid]
        if qids:
            return sorted(set(qids), key=natural_q_key)
    model_route = load_json(OUTPUT_DIR / "plan" / "model_route.json")
    qids = []
    for item in model_route.get("questions", []) if isinstance(model_route, dict) else []:
        if isinstance(item, dict) and item.get("question_id"):
            qids.append(str(item["question_id"]))
    return sorted(set(qids), key=natural_q_key)


def index_items(data: Any, key: str, id_key: str) -> list[dict[str, Any]]:
    if not isinstance(data, dict):
        return []
    return [item for item in data.get(key, []) if isinstance(item, dict) and item.get(id_key)]


def referenced(text: str, item: dict[str, Any], id_key: str) -> bool:
    candidates = [
        str(item.get(id_key) or ""),
        str(item.get("title") or ""),
        Path(str(item.get("path") or item.get("expected_path") or "")).stem,
    ]
    candidates = [candidate.strip() for candidate in candidates if candidate and candidate.strip()]
    return any(candidate in text for candidate in candidates)


def check_docx_structure(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {"exists": False}
    try:
        doc = Document(str(path))
        headings = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.style and paragraph.style.name.startswith("Heading") and paragraph.text.strip()
        ]
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        return {
            "exists": True,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "image_count": image_count,
            "heading_count": len(headings),
            "sample_headings": headings[:12],
        }
    except Exception as exc:
        return {"exists": True, "error": str(exc)}


def markdown_heading_count(text: str) -> int:
    return len(re.findall(r"(^|\n)\s*#{1,6}\s+\S+", text))


def visual_qa_failures(
    docx_structure: dict[str, Any],
    source_heading_count: int,
    figure_count: int,
    table_count: int,
) -> tuple[list[str], list[str]]:
    failures: list[str] = []
    warnings: list[str] = []
    if not docx_structure.get("exists") or docx_structure.get("error"):
        return failures, warnings

    paragraph_count = int(docx_structure.get("paragraph_count") or 0)
    heading_count = int(docx_structure.get("heading_count") or 0)
    docx_table_count = int(docx_structure.get("table_count") or 0)
    image_count = int(docx_structure.get("image_count") or 0)

    if paragraph_count < 10:
        failures.append(f"Word 段落数量异常偏少：{paragraph_count} < 10")
    if source_heading_count > 0 and heading_count == 0:
        failures.append("Word 中没有可识别标题样式，标题结构可能未正确写入。")
    elif source_heading_count > 0 and heading_count < max(1, source_heading_count // 2):
        warnings.append(f"Word 标题数量明显少于 Markdown 标题：{heading_count} < {source_heading_count}")

    if figure_count > 0 and image_count == 0:
        failures.append("figure_index.json 有图片计划，但 Word 中没有图片。")
    elif image_count < figure_count:
        warnings.append(f"Word 图片数量少于 figure_index.json：{image_count} < {figure_count}")

    if table_count > 0 and docx_table_count == 0:
        failures.append("table_index.json 有表格计划，但 Word 中没有表格。")
    elif docx_table_count < table_count:
        warnings.append(f"Word 表格数量少于 table_index.json：{docx_table_count} < {table_count}")

    return failures, warnings


def evaluate() -> dict[str, Any]:
    source = source_path()
    outline = load_json(OUTLINE_FILE)
    figure_index = load_json(FIGURE_INDEX_FILE)
    table_index = load_json(TABLE_INDEX_FILE)
    failures: list[str] = []
    warnings: list[str] = []

    if not source.exists():
        failures.append(f"缺少正式论文源文件：{rel(SOURCE_FILE)}")
        text = ""
    else:
        text = source.read_text(encoding="utf-8")

    counts = char_count(text)
    target_words = outline.get("target_words", {}) if isinstance(outline, dict) else {}
    min_words = int(target_words.get("min", 18000) or 18000)
    max_words = int(target_words.get("max", 25000) or 25000)
    if counts["content"] < min_words:
        failures.append(f"正文有效字数不足：{counts['content']} < {min_words}")
    if counts["content"] > max_words:
        warnings.append(f"正文有效字数超过建议上限：{counts['content']} > {max_words}")

    missing_sections = [label for label in REQUIRED_SECTIONS if not has_required_section(text, label)]
    for label in missing_sections:
        failures.append(f"缺少正式论文结构：{label}")

    if not re.search(r"(^|\n)\s*#*\s*5\.1\s+", text):
        failures.append("缺少 5.1 问题一模型章节。")
    if not re.search(r"(^|\n)\s*#*\s*5\.1\.1\s+", text):
        failures.append("缺少 5.1.1 三级标题。")
    if not re.search(r"(^|\n)\s*#*\s*5\.1\.2\s+", text):
        failures.append("缺少 5.1.2 三级标题。")

    question_reports = []
    for index, qid in enumerate(qids_from_outline(outline), start=1):
        q_failures: list[str] = []
        section = f"5.{index}"
        section_pattern = rf"(^|\n)\s*#*\s*{re.escape(section)}\s+"
        if not re.search(section_pattern, text):
            q_failures.append(f"缺少 {section} 对应 {qid} 的模型章节")
        for suffix, title in (
            ("1", "建模思路"),
            ("2", "变量定义与公式推导"),
            ("3", "求解算法"),
            ("4", "结果分析"),
            ("5", "模型检验或灵敏度分析"),
        ):
            if not re.search(rf"(^|\n)\s*#*\s*{re.escape(section + '.' + suffix)}\s+.*{title}", text):
                q_failures.append(f"缺少 {section}.{suffix} {title}")
        if not re.search(rf"{qid}|问题[一二三四五六七八九十{index}]", text):
            q_failures.append(f"正文未明确回扣 {qid}")
        if not re.search(rf"{section}[\s\S]{{0,5000}}Step\s*1", text, flags=re.IGNORECASE):
            q_failures.append(f"{section} 缺少 Step 1/Step 2 形式的算法步骤")
        question_reports.append({"question_id": qid, "status": "FAIL" if q_failures else "PASS", "failures": q_failures})
        failures.extend(q_failures)

    figures = index_items(figure_index, "figures", "figure_id")
    tables = index_items(table_index, "tables", "table_id")
    missing_figures = [item.get("figure_id") for item in figures if not referenced(text, item, "figure_id")]
    missing_tables = [item.get("table_id") for item in tables if not referenced(text, item, "table_id")]
    for figure_id in missing_figures:
        failures.append(f"figure_index.json 中的图片未在正文引用：{figure_id}")
    for table_id in missing_tables:
        failures.append(f"table_index.json 中的表格未在正文引用：{table_id}")
    if len(figures) < 5:
        warnings.append(f"图数量少于展示样例建议值：{len(figures)} < 5")
    if len(tables) < 5:
        warnings.append(f"表数量少于展示样例建议值：{len(tables)} < 5")

    for placeholder in PLACEHOLDERS:
        if placeholder in text:
            failures.append(f"存在占位符或待补文本：{placeholder}")

    if "参考文献" in text and len(re.findall(r"\[\d+\]", text)) < 3:
        warnings.append("参考文献条目少于 3 条，建议补充权威来源。")

    docx_structure = check_docx_structure(DOCX_FILE)
    if not docx_structure.get("exists"):
        failures.append(f"缺少正式 Word 文件：{rel(DOCX_FILE)}")
    elif docx_structure.get("error"):
        failures.append(f"Word 文件无法读取：{docx_structure['error']}")

    source_heading_count = markdown_heading_count(text)
    visual_failures, visual_warnings = visual_qa_failures(docx_structure, source_heading_count, len(figures), len(tables))
    failures.extend(visual_failures)
    warnings.extend(visual_warnings)

    return {
        "schema_version": "1.0",
        "generated_by": "paper-formal-writer/scripts/check_paper_format.py",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "status": "PASS" if not failures else "FAIL",
        "source": rel(source),
        "docx": rel(DOCX_FILE),
        "counts": counts,
        "target_words": {"min": min_words, "max": max_words},
        "question_reports": question_reports,
        "figure_count": len(figures),
        "table_count": len(tables),
        "missing_figures": missing_figures,
        "missing_tables": missing_tables,
        "source_heading_count": source_heading_count,
        "docx_structure": docx_structure,
        "visual_qa": {
            "status": "PASS" if not visual_failures else "FAIL",
            "failures": visual_failures,
            "warnings": visual_warnings,
        },
        "failures": failures,
        "warnings": warnings,
    }


def write_reports(report: dict[str, Any]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# Formal Paper Format Check Report",
        "",
        f"- Status: `{report['status']}`",
        f"- Generated at: `{report['generated_at']}`",
        f"- Source: `{report['source']}`",
        f"- DOCX: `{report['docx']}`",
        f"- Effective chars: `{report['counts']['content']}`",
        f"- CJK chars: `{report['counts']['cjk']}`",
        f"- Figures in index: `{report['figure_count']}`",
        f"- Tables in index: `{report['table_count']}`",
        "",
    ]
    if report["failures"]:
        lines.append("## Failures")
        lines.extend(f"- {item}" for item in report["failures"])
        lines.append("")
    if report["warnings"]:
        lines.append("## Warnings")
        lines.extend(f"- {item}" for item in report["warnings"])
        lines.append("")
    lines.append("## Question Coverage")
    for item in report["question_reports"]:
        lines.append(f"- {item['question_id']}: `{item['status']}`")
        for failure in item["failures"]:
            lines.append(f"  - {failure}")
    lines.append("")
    lines.append("## DOCX Structure")
    for key, value in report["docx_structure"].items():
        lines.append(f"- {key}: `{value}`")
    lines.append("")
    lines.append("## Visual QA")
    lines.append(f"- status: `{report['visual_qa']['status']}`")
    lines.append(f"- source_heading_count: `{report['source_heading_count']}`")
    for failure in report["visual_qa"]["failures"]:
        lines.append(f"- failure: {failure}")
    for warning in report["visual_qa"]["warnings"]:
        lines.append(f"- warning: {warning}")
    REPORT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> int:
    configure_utf8_stdio()
    report = evaluate()
    write_reports(report)
    print(f"正式论文格式检查报告：{rel(REPORT_MD)}")
    if report["status"] == "PASS":
        print("✅ 正式论文格式门禁通过。")
        return 0
    print("⚠️ 正式论文格式门禁未通过。")
    for failure in report["failures"][:12]:
        print(f" - {failure}")
    if len(report["failures"]) > 12:
        print(f" - 其余 {len(report['failures']) - 12} 项见报告。")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
