import json
import re
from pathlib import Path
from typing import Dict, List

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


BASE_DIR = Path.cwd()
OUTPUT_DIR = BASE_DIR / "paper_output"
TASKS_FILE = OUTPUT_DIR / "tasks.json"
UNITS_DIR = OUTPUT_DIR / "micro_units"
FINAL_FILE = OUTPUT_DIR / "final_paper.md"
REF_REPORT = OUTPUT_DIR / "ref_check.md"


def collect_units_from_tasks() -> List[str]:
    if not TASKS_FILE.exists():
        files = sorted(UNITS_DIR.glob("*.txt"), key=lambda x: x.stem)
        return [f.read_text(encoding="utf-8").strip() for f in files]

    tasks = json.loads(TASKS_FILE.read_text(encoding="utf-8"))
    units: List[str] = []
    last_section = None
    for t in tasks:
        section = t.get("section")
        if section and section != last_section:
            units.append(f"## {section}\n")
            last_section = section
        fp = Path(t.get("file_path", ""))
        if fp.exists():
            units.append(fp.read_text(encoding="utf-8").strip())
        else:
            unit_id = t.get("id", "")
            units.append(f"【{unit_id}】\n(缺失)\n")
    return units


def auto_numbering(text: str) -> str:
    fig_counter = 1

    def fig_repl(_m):
        nonlocal fig_counter
        s = f"图{fig_counter}"
        fig_counter += 1
        return s

    text = re.sub(r"图\s*\d+", fig_repl, text)

    tab_counter = 1

    def tab_repl(_m):
        nonlocal tab_counter
        s = f"表{tab_counter}"
        tab_counter += 1
        return s

    text = re.sub(r"表\s*\d+", tab_repl, text)

    eq_counter = 1

    def eq_repl(_m):
        nonlocal eq_counter
        s = f"式({eq_counter})"
        eq_counter += 1
        return s

    text = re.sub(r"式\s*\(\s*\d+\s*\)", eq_repl, text)

    ref_counter = 1

    def ref_repl(_m):
        nonlocal ref_counter
        s = f"[{ref_counter}]"
        ref_counter += 1
        return s

    text = re.sub(r"\[\s*\d+\s*\]", ref_repl, text)
    return text


def build_cross_ref_index(text: str) -> Dict[str, str]:
    index: Dict[str, str] = {}
    for m in re.finditer(r"图\s*(\d+)", text):
        k = f"图{m.group(1)}"
        index[k] = k
    for m in re.finditer(r"表\s*(\d+)", text):
        k = f"表{m.group(1)}"
        index[k] = k
    for m in re.finditer(r"式\s*\(\s*(\d+)\s*\)", text):
        k = f"式({m.group(1)})"
        index[k] = k
    return index


def replace_cross_ref(text: str, index: Dict[str, str]) -> str:
    for old, new in index.items():
        text = text.replace(f"见{old}", f"见{new}")
    return text


def generate_toc(text: str) -> str:
    toc_lines = ["# 目录\n"]
    h2 = re.findall(r"^##\s+(.+)$", text, flags=re.MULTILINE)
    if not h2:
        return ""
    for i, name in enumerate(h2, start=1):
        toc_lines.append(f"{i}. {name.strip()}")
    return "\n".join(toc_lines) + "\n\n"


def ref_check_report(text: str) -> str:
    lines = ["# 引用断链检查报告\n"]

    defined_fig = {f"图{m.group(1)}" for m in re.finditer(r"^图\s*(\d+)", text, flags=re.MULTILINE)}
    defined_tab = {f"表{m.group(1)}" for m in re.finditer(r"^表\s*(\d+)", text, flags=re.MULTILINE)}
    defined_eq = {f"式({m.group(1)})" for m in re.finditer(r"^式\(\s*(\d+)\s*\)", text, flags=re.MULTILINE)}

    for m in re.finditer(r"见式\s*\(\s*(\d+)\s*\)", text):
        eq_id = f"式({m.group(1)})"
        if eq_id not in defined_eq:
            lines.append(f"- 断链：{eq_id} 被引用但未定义")
    for m in re.finditer(r"见图\s*(\d+)", text):
        fig_id = f"图{m.group(1)}"
        if fig_id not in defined_fig:
            lines.append(f"- 断链：{fig_id} 被引用但未定义")
    for m in re.finditer(r"见表\s*(\d+)", text):
        tab_id = f"表{m.group(1)}"
        if tab_id not in defined_tab:
            lines.append(f"- 断链：{tab_id} 被引用但未定义")
    if len(lines) == 1:
        lines.append("- 未检测到断链")
    return "\n".join(lines) + "\n"


class SimpleMarkdownToDocx:
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx not installed")
        self.doc = Document()
        
    def convert(self, md_text: str, output_path: Path):
        self.doc.add_heading('数学建模论文 (自动生成)', 0)
        
        # Warning note
        p = self.doc.add_paragraph()
        run = p.add_run('注意：本 Word 文档由脚本直接生成（未使用 Pandoc），数学公式将显示为 LaTeX 源码。')
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        def sanitize(s):
            # Remove control characters that are not allowed in XML
            return re.sub(r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]', '', s)

        # Simple line-based parser
        for line in md_text.split('\n'):
            line = sanitize(line.strip())
            if not line:
                continue
            
            # Headings
            if line.startswith('#'):
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                text = line[level:].strip()
                # docx headings 1-9
                if 1 <= level <= 9:
                    self.doc.add_heading(text, level=level)
                else:
                    self.doc.add_paragraph(line)
                continue
            
            # Images: ![alt](path)
            # Regex to handle image at start of line
            img_match = re.match(r'^!\[.*?\]\((.*?)\)$', line)
            if img_match:
                img_rel_path = img_match.group(1)
                
                # Try to resolve path
                img_path = Path(img_rel_path)
                if not img_path.exists():
                     # Check relative to OUTPUT_DIR
                     img_path = OUTPUT_DIR / img_rel_path
                
                if not img_path.exists():
                    # Check relative to root
                     img_path = BASE_DIR / img_rel_path
                     
                if img_path.exists():
                    try:
                        self.doc.add_picture(str(img_path), width=Inches(6))
                        # Center the image
                        if self.doc.paragraphs:
                            last_p = self.doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        self.doc.add_paragraph(f"[图片加载失败: {img_rel_path}]")
                else:
                    self.doc.add_paragraph(f"[图片未找到: {img_rel_path}]")
                continue

            # Default paragraph
            self.doc.add_paragraph(line)
            
        self.doc.save(output_path)


def main() -> None:
    units = collect_units_from_tasks()
    raw_text = "\n\n".join(units)
    numbered = auto_numbering(raw_text)
    index = build_cross_ref_index(numbered)
    final = replace_cross_ref(numbered, index)
    toc = generate_toc(final)
    full = toc + final
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FINAL_FILE.write_text(full, encoding="utf-8")
    REF_REPORT.write_text(ref_check_report(full), encoding="utf-8")
    print(f"合并完成：{FINAL_FILE}")
    print(f"引用检查：{REF_REPORT}")
    
    # Export to Docx directly
    if HAS_DOCX:
        docx_path = OUTPUT_DIR / "final_paper_direct.docx"
        try:
            converter = SimpleMarkdownToDocx()
            converter.convert(full, docx_path)
            print(f"Word导出成功：{docx_path}")
        except Exception as e:
            print(f"Word导出失败：{e}")
    else:
        print("跳过Word导出：未安装 python-docx")


if __name__ == "__main__":
    main()
