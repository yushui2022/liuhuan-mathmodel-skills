import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any


BASE_DIR = Path.cwd()
PROBLEM_DIR = BASE_DIR / "problem_files"
OUTPUT_DIR = BASE_DIR / "paper_output"
STEP1_DIR = OUTPUT_DIR / "step1"
ANALYSIS_FILE = STEP1_DIR / "problem_analysis.json"
DATA_REQUIREMENTS_FILE = BASE_DIR / "data_requirements.json"


QUESTION_NUMERALS = {
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
}


MODEL_RULES = [
    {
        "task_type": "预测/回归",
        "keywords": ["预测", "预报", "趋势", "估计", "回归", "未来", "时间序列", "forecast"],
        "baseline": "线性回归、移动平均、指数平滑或 ARIMA",
        "improved": "特征工程 + 树模型、集成学习或时序深度模型",
        "validation": ["留出验证或滚动回测", "RMSE/MAE/MAPE", "残差分析与置信区间"],
        "figures": ["真实值-预测值对比图", "残差分布图", "趋势外推图"],
    },
    {
        "task_type": "优化/规划",
        "keywords": ["优化", "最优", "最大", "最小", "分配", "调度", "路径", "选址", "成本", "收益", "约束"],
        "baseline": "线性规划、整数规划或可行规则基线",
        "improved": "多目标规划、启发式算法、遗传算法或模拟退火",
        "validation": ["可行性检查", "约束满足率", "与基准策略对比", "敏感性分析"],
        "figures": ["方案对比柱状图", "约束满足情况图", "敏感性分析折线图"],
    },
    {
        "task_type": "评价/排序",
        "keywords": ["评价", "评估", "排序", "排名", "指标体系", "权重", "综合", "打分", "优劣"],
        "baseline": "归一化 + 加权求和",
        "improved": "熵权法、CRITIC、AHP、TOPSIS 或 VIKOR",
        "validation": ["权重敏感性分析", "排名稳定性检验", "与已知事实对照"],
        "figures": ["指标权重图", "综合得分排序图", "雷达图或热力图"],
    },
    {
        "task_type": "分类/识别",
        "keywords": ["分类", "识别", "判别", "类别", "风险等级", "是否", "检测"],
        "baseline": "逻辑回归、朴素贝叶斯或决策树",
        "improved": "随机森林、梯度提升或代价敏感学习",
        "validation": ["混淆矩阵", "Precision/Recall/F1", "ROC/AUC 或 PR 曲线"],
        "figures": ["混淆矩阵热力图", "ROC 曲线", "特征重要性图"],
    },
    {
        "task_type": "聚类/分群",
        "keywords": ["聚类", "分群", "画像", "相似性", "模式发现", "群体"],
        "baseline": "K-means 或层次聚类",
        "improved": "GMM、DBSCAN 或谱聚类",
        "validation": ["轮廓系数", "聚类稳定性", "群体差异解释"],
        "figures": ["降维散点图", "群体特征对比图", "聚类热力图"],
    },
    {
        "task_type": "机理/仿真",
        "keywords": ["机理", "仿真", "微分", "动力学", "物理", "传播", "过程", "系统"],
        "baseline": "差分方程、微分方程或机理方程",
        "improved": "参数校准、不确定性分析或数据驱动校正",
        "validation": ["历史拟合", "参数敏感性", "极端情景测试"],
        "figures": ["仿真曲线图", "参数敏感性图", "情景对比图"],
    },
]


def safe_read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return path.read_text(encoding=encoding, errors="ignore")
        except Exception:
            continue
    return ""


def read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        doc = Document(path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""


def read_pdf(path: Path) -> str:
    try:
        import fitz  # type: ignore

        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    except Exception:
        pass
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        pass
    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        return ""


def collect_problem_files() -> list[Path]:
    if not PROBLEM_DIR.exists():
        return []
    return sorted([p for p in PROBLEM_DIR.rglob("*") if p.is_file()], key=lambda p: str(p).lower())


def extract_text(files: list[Path]) -> tuple[str, list[dict[str, Any]], list[dict[str, Any]]]:
    text_parts: list[str] = []
    documents: list[dict[str, Any]] = []
    data_files: list[dict[str, Any]] = []

    for path in files:
        suffix = path.suffix.lower()
        rel = str(path.relative_to(BASE_DIR)) if path.is_relative_to(BASE_DIR) else str(path)
        content = ""

        if suffix in {".txt", ".md"}:
            content = safe_read_text(path)
            documents.append({"path": rel, "type": suffix.lstrip("."), "chars": len(content)})
        elif suffix == ".docx":
            content = read_docx(path)
            documents.append({"path": rel, "type": "docx", "chars": len(content)})
        elif suffix == ".pdf":
            content = read_pdf(path)
            documents.append({"path": rel, "type": "pdf", "chars": len(content)})
        elif suffix in {".csv", ".xlsx", ".xls"}:
            data_files.append(profile_data_file(path))

        if content.strip():
            text_parts.append(f"\n\n# 文件：{rel}\n{content.strip()}")

    return "\n".join(text_parts).strip(), documents, data_files


def profile_data_file(path: Path) -> dict[str, Any]:
    rel = str(path.relative_to(BASE_DIR)) if path.is_relative_to(BASE_DIR) else str(path)
    profile: dict[str, Any] = {"path": rel, "type": path.suffix.lower().lstrip("."), "readable": False}
    try:
        import pandas as pd

        if path.suffix.lower() == ".csv":
            df = pd.read_csv(path, nrows=20)
            profile.update({"readable": True, "rows_sampled": len(df), "columns": [str(c) for c in df.columns]})
        elif path.suffix.lower() in {".xlsx", ".xls"}:
            xls = pd.ExcelFile(path)
            sheets = []
            for sheet in xls.sheet_names[:5]:
                df = pd.read_excel(path, sheet_name=sheet, nrows=20)
                sheets.append({"sheet": sheet, "rows_sampled": len(df), "columns": [str(c) for c in df.columns]})
            profile.update({"readable": True, "sheets": sheets})
    except Exception as exc:
        profile["error"] = str(exc)
    return profile


def normalize_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chinese_num_to_int(value: str) -> int | None:
    value = value.strip()
    if value.isdigit():
        return int(value)
    reverse = {v: k for k, v in QUESTION_NUMERALS.items()}
    return reverse.get(value)


def split_questions(text: str) -> list[dict[str, str]]:
    if not text:
        return []

    pattern = re.compile(r"(?:^|\n)\s*(?:问题|任务|第)\s*([一二三四五六七八九十\d]+)\s*(?:问|题)?[：:、.．\s]*(.*)")
    matches = list(pattern.finditer(text))
    questions: list[dict[str, str]] = []

    for idx, match in enumerate(matches):
        num = chinese_num_to_int(match.group(1)) or (idx + 1)
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        chunk = text[start:end].strip()
        first_line = match.group(2).strip() or f"问题{QUESTION_NUMERALS.get(num, num)}"
        questions.append(
            {
                "id": f"Q{num}",
                "title": f"问题{QUESTION_NUMERALS.get(num, num)}",
                "summary": first_line[:160],
                "raw_text": chunk[:2000],
            }
        )

    if questions:
        unique: dict[str, dict[str, str]] = {}
        for item in questions:
            unique.setdefault(item["id"], item)
        return list(unique.values())

    return [
        {
            "id": "Q1",
            "title": "问题一",
            "summary": "根据赛题文本和附件数据完成核心建模任务。",
            "raw_text": text[:2000],
        }
    ]


def choose_model(text: str) -> dict[str, Any]:
    lower = text.lower()
    best = None
    best_score = -1
    for rule in MODEL_RULES:
        score = sum(1 for kw in rule["keywords"] if kw.lower() in lower)
        if score > best_score:
            best = rule
            best_score = score
    if best is None or best_score <= 0:
        best = {
            "task_type": "综合建模/统计分析",
            "baseline": "描述统计、相关性分析或可解释基线模型",
            "improved": "结合题目目标选择回归、评价、优化或仿真模型",
            "validation": ["结果复核", "敏感性分析", "与题意约束对照"],
            "figures": ["数据概览图", "结果对比图", "敏感性分析图"],
        }
    return {
        "task_type": best["task_type"],
        "baseline_model": best["baseline"],
        "improved_model": best["improved"],
        "validation_plan": best["validation"],
        "figure_suggestions": best["figures"],
    }


def extract_constraints(text: str) -> list[str]:
    sentences = re.split(r"[。；;\n]", text)
    keywords = ["要求", "约束", "限制", "必须", "至少", "不超过", "不能", "需要", "给出", "确定", "建立"]
    hits = []
    for sent in sentences:
        sent = sent.strip()
        if 8 <= len(sent) <= 180 and any(kw in sent for kw in keywords):
            hits.append(sent)
        if len(hits) >= 8:
            break
    return hits


def detect_external_data_need(text: str) -> list[dict[str, Any]]:
    keywords = ["搜集", "收集", "查找", "公开数据", "外部数据", "统计数据", "权威数据", "网络数据"]
    if not any(kw in text for kw in keywords):
        return []
    query = " ".join(extract_constraints(text)[:2]) or "根据赛题补充权威公开数据"
    return [
        {
            "type": "manual_search",
            "query": query[:200],
            "notes": "由 problem-doc-model-selector 检测到题目可能需要外部公开数据，请优先寻找官方或权威来源。",
            "active": True,
        }
    ]


def build_analysis(text: str, documents: list[dict[str, Any]], data_files: list[dict[str, Any]]) -> dict[str, Any]:
    text = normalize_text(text)
    constraints = extract_constraints(text)
    raw_questions = split_questions(text)
    questions = []

    for item in raw_questions:
        model = choose_model(item["raw_text"] + "\n" + item["summary"])
        questions.append(
            {
                **item,
                "task_type": model["task_type"],
                "inputs": ["赛题文本", "附件数据", "必要的外部权威数据"],
                "outputs": ["可量化结果", "图表证据", "对应原问的结论"],
                "constraints": extract_constraints(item["raw_text"]) or constraints[:3],
                "recommended_models": {
                    "baseline": model["baseline_model"],
                    "improved": model["improved_model"],
                },
                "validation_plan": model["validation_plan"],
                "figure_suggestions": model["figure_suggestions"],
            }
        )

    data_requirements = detect_external_data_need(text)
    return {
        "schema_version": "1.0",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "generated_by": "problem-doc-model-selector/scripts/analyze_problem.py",
        "source_dir": str(PROBLEM_DIR),
        "documents": documents,
        "data_files": data_files,
        "problem_text_excerpt": text[:3000],
        "global_constraints": constraints,
        "questions": questions,
        "data_requirements": data_requirements,
        "next_outputs": {
            "model_route": "paper_output/plan/model_route.json",
            "tasks": "paper_output/tasks.json",
            "figures": "paper_output/figures/",
        },
    }


def write_markdown_outputs(analysis: dict[str, Any]) -> None:
    questions = analysis.get("questions", [])
    lines_a = ["# A 题意对齐\n"]
    lines_c = ["# C 评分点对齐表\n", "| 评分点 | 证据形式 | 论文位置 |\n|---|---|---|\n"]
    outline = [
        "# B 论文大纲\n",
        "1. 摘要\n",
        "2. 问题重述\n",
        "3. 模型假设与符号说明\n",
        "4. 数据说明与预处理\n",
    ]
    route: list[dict[str, Any]] = []

    for idx, q in enumerate(questions, start=1):
        lines_a.append(f"## {q['title']}\n")
        lines_a.append(f"- 任务类型：{q['task_type']}\n")
        lines_a.append(f"- 题意摘要：{q.get('summary', '')}\n")
        lines_a.append(f"- 基线模型：{q['recommended_models']['baseline']}\n")
        lines_a.append(f"- 改进路线：{q['recommended_models']['improved']}\n")
        lines_a.append(f"- 验证计划：{'；'.join(q.get('validation_plan', []))}\n")
        lines_a.append(f"- 建议图表：{'；'.join(q.get('figure_suggestions', []))}\n\n")

        outline.append(f"{4 + idx}. {q['title']}：模型建立、求解、结果与检验\n")
        lines_c.append(
            f"| {q['title']}回答完整 | 模型公式、结果表、验证图 | {q['title']}模型与结果分析 |\n"
        )
        route.append(
            {
                "question_id": q["id"],
                "title": q["title"],
                "task_type": q["task_type"],
                "baseline_model": q["recommended_models"]["baseline"],
                "improved_model": q["recommended_models"]["improved"],
                "validation_plan": q.get("validation_plan", []),
                "figure_suggestions": q.get("figure_suggestions", []),
            }
        )

    outline.extend(
        [
            f"{5 + len(questions)}. 综合结果分析与敏感性检验\n",
            f"{6 + len(questions)}. 模型评价、推广与不足\n",
            f"{7 + len(questions)}. 结论、参考文献与附录\n",
        ]
    )
    lines_c.extend(
        [
            "| 数据可复现 | 清洗脚本、字段说明、来源记录 | 数据说明与附录 |\n",
            "| 图文一致 | 图表文件与正文引用 | 结果分析 |\n",
            "| 结论对应原问 | 分问题结论清单 | 结论 |\n",
        ]
    )

    STEP1_DIR.mkdir(parents=True, exist_ok=True)
    (STEP1_DIR / "A_题意对齐.md").write_text("".join(lines_a), encoding="utf-8")
    (STEP1_DIR / "B_论文大纲.md").write_text("".join(outline), encoding="utf-8")
    (STEP1_DIR / "C_评分点对齐表.md").write_text("".join(lines_c), encoding="utf-8")
    (STEP1_DIR / "D_模型路线.json").write_text(json.dumps(route, ensure_ascii=False, indent=2), encoding="utf-8")


def write_data_requirements(analysis: dict[str, Any]) -> None:
    tasks = analysis.get("data_requirements") or []
    if not tasks:
        return
    DATA_REQUIREMENTS_FILE.write_text(json.dumps({"tasks": tasks}, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    STEP1_DIR.mkdir(parents=True, exist_ok=True)

    files = collect_problem_files()
    if not files:
        print(f"⚠️ 未找到赛题或附件目录：{PROBLEM_DIR}")
        return 1

    text, documents, data_files = extract_text(files)
    analysis = build_analysis(text, documents, data_files)
    ANALYSIS_FILE.write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown_outputs(analysis)
    write_data_requirements(analysis)

    print(f"✅ 已生成结构化赛题分析：{ANALYSIS_FILE}")
    print(f"   子问题数量：{len(analysis.get('questions', []))}")
    print(f"   文档数量：{len(documents)}，数据附件数量：{len(data_files)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
